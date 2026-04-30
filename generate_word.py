#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成喵了个鱼参赛总结Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def set_chinese_font(run, font_name='Microsoft YaHei', font_size=10, bold=False, color=None):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)

def generate_word():
    # 创建文档
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 读取Markdown内容
    with open('喵了个鱼_参赛总结.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 移除emoji
    content = re.sub(r'[\U00010000-\U0010ffff]', '', content)
    
    # 添加标题
    title = doc.add_heading('', level=0)
    title_run = title.add_run('喵了个鱼 - Trae编程比赛参赛总结')
    set_chinese_font(title_run, font_size=24, bold=True, color=(255, 107, 107))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('微信小程序开发案例')
    set_chinese_font(subtitle_run, font_size=14, color=(90, 74, 66))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 按行处理内容
    lines = content.split('\n')
    i = 0
    code_buffer = []
    in_code_block = False
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_buffer:
                    # 添加代码块
                    code_para = doc.add_paragraph()
                    code_text = '\n'.join(code_buffer)
                    code_run = code_para.add_run(code_text)
                    set_chinese_font(code_run, font_name='Courier New', font_size=9)
                    code_para.paragraph_format.left_indent = Inches(0.3)
                    code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            h1 = doc.add_heading('', level=1)
            h1_run = h1.add_run(line[2:])
            set_chinese_font(h1_run, font_size=18, bold=True, color=(255, 107, 107))
        elif line.startswith('## '):
            h2 = doc.add_heading('', level=2)
            h2_run = h2.add_run(line[3:])
            set_chinese_font(h2_run, font_size=14, bold=True, color=(90, 74, 66))
        elif line.startswith('### '):
            h3 = doc.add_heading('', level=3)
            h3_run = h3.add_run(line[4:])
            set_chinese_font(h3_run, font_size=12, bold=True, color=(90, 74, 66))
        elif line.startswith('#### '):
            h4 = doc.add_heading('', level=4)
            h4_run = h4.add_run(line[5:])
            set_chinese_font(h4_run, font_size=11, bold=True, color=(90, 74, 66))
        # 处理引用
        elif line.startswith('>'):
            quote_para = doc.add_paragraph()
            quote_run = quote_para.add_run(line[1:].strip())
            set_chinese_font(quote_run, font_size=10, color=(102, 102, 102))
            quote_para.paragraph_format.left_indent = Inches(0.3)
        # 处理表格行（简化处理）
        elif line.startswith('|'):
            # 跳过表格分隔线
            if '---' in line:
                i += 1
                continue
            # 提取表格内容
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                table_para = doc.add_paragraph()
                table_run = table_para.add_run(' | '.join(cells))
                set_chinese_font(table_run, font_size=9)
                table_para.paragraph_format.left_indent = Inches(0.2)
        # 处理分隔线
        elif line == '---':
            doc.add_paragraph('_' * 50)
        # 处理列表项
        elif line.startswith('- ') or line.startswith('* '):
            list_para = doc.add_paragraph(style='List Bullet')
            list_run = list_para.add_run(line[2:])
            set_chinese_font(list_run, font_size=10)
        elif re.match(r'^\d+\. ', line):
            list_para = doc.add_paragraph(style='List Number')
            list_run = list_para.add_run(re.sub(r'^\d+\. ', '', line))
            set_chinese_font(list_run, font_size=10)
        # 处理普通段落
        elif line:
            # 处理Markdown标记
            para = doc.add_paragraph()
            
            # 分割文本处理粗体和斜体
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|`.*?`)', line)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # 粗体
                    run = para.add_run(part[2:-2])
                    set_chinese_font(run, font_size=10, bold=True)
                elif part.startswith('__') and part.endswith('__'):
                    # 粗体
                    run = para.add_run(part[2:-2])
                    set_chinese_font(run, font_size=10, bold=True)
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    # 斜体
                    run = para.add_run(part[1:-1])
                    set_chinese_font(run, font_size=10)
                    run.font.italic = True
                elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                    # 斜体
                    run = para.add_run(part[1:-1])
                    set_chinese_font(run, font_size=10)
                    run.font.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    # 代码
                    run = para.add_run(part[1:-1])
                    set_chinese_font(run, font_name='Courier New', font_size=9)
                    run.font.color.rgb = RGBColor(211, 51, 132)
                else:
                    # 普通文本
                    if part:
                        run = para.add_run(part)
                        set_chinese_font(run, font_size=10)
        else:
            # 空行
            doc.add_paragraph()
        
        i += 1
    
    # 保存文档
    output_file = '喵了个鱼_Trae编程比赛参赛总结.docx'
    doc.save(output_file)
    print(f"✅ Word文档生成成功！")
    print(f"📄 文件名：{output_file}")
    print(f"📍 位置：/Users/ouyangguoqing/Documents/trae_projects/miaolegeyu1/{output_file}")

if __name__ == '__main__':
    generate_word()
